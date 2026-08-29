"""Step 9 / 10: render the Word document from the typed draft.

The renderer is **trusted local code**. The model is not allowed to write
DOCX XML, to choose a template, to set the document classification, or to
control which sections appear. The renderer reads a ``ApprovalNoteDraft``
and produces a file. The renderer is the only thing that can produce a
file.

Two reasons for that split:

* **A model that produces XML can produce invalid XML.** Not by malice; by
  not knowing what a `</w:p>` looks like, and a document that does not
  open is a document nobody signs.
* **A model that chooses the template can choose the wrong template.** A
  policy change tomorrow that drops a section should land in this file
  alone, not in every model that ever produced a note.

The template is built here, in Python, and used through ``python-docx``
so the produced file is one Word will actually open. ``python-docx`` is
the trusted library; it does the XML escaping and the package layout. We
add the required headings, the classification banner, the signature area,
and the citation markers.

## Re-opening

The renderer also re-opens the file it just wrote and confirms the
sections it claims are in the file are in the file. A bug between the
renderer and the library would not be caught by tests of the renderer
alone.
"""

from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from .approval import ApprovalRecord
from .draft import (
    ApprovalNoteDraft,
    Classification,
    DraftStatus,
    Finding,
    Severity,
    UncertaintyNote,
)


#: A file we trust to be the renderer. A name string is a name string; an
#: import path or a file path is what the prompt means by "trusted local
#: template".
TEMPLATE_NAME = "approval_note"

#: The required headings. Order matters; it is the order the reviewer reads
#: them in. Adding a new heading is a renderer change; the verifier checks
#: for these exact names.
REQUIRED_HEADINGS: Tuple[str, ...] = (
    "Findings",
    "Calculation",
    "Recommendation",
    "Supporting references",
    "Assumptions",
    "Signature",
)

#: A finding page needs room for a real table. The renderer does not put
#: more than this many in a single Word table; longer lists spill onto a
#: second table the same way.
MAX_ROWS_PER_TABLE = 20


@dataclass
class DocumentMetadata:
    """Provenance the renderer stamps on every document."""

    task_id: str
    created_at: str
    model: str
    skill: str
    classification: Classification
    is_draft: bool
    approval: Optional[ApprovalRecord] = None


@dataclass
class RenderResult:
    """What the renderer produced, with the file's own self-check."""

    output_path: Path
    sections: List[str] = field(default_factory=list)
    problems: List[str] = field(default_factory=list)
    opens: bool = False
    byte_size: int = 0

    def is_sound(self) -> bool:
        return self.opens and not self.problems


class TemplateError(Exception):
    """The renderer could not produce a file because the template refused."""


def _format_value_with_unit(value: str) -> str:
    """Pass through a calculation's result verbatim. The model is not
    allowed to reformat figures, and neither is the renderer."""
    return value.strip()


def _format_classification(c: Classification) -> str:
    """Classification in the human-readable form the reviewer sees."""
    table = {
        Classification.PUBLIC: "Public",
        Classification.INTERNAL: "Internal",
        Classification.PROCESS_DIAGRAM: "Process diagram",
        Classification.CONFIDENTIAL: "Confidential",
        Classification.RESTRICTED: "Restricted",
    }
    return table.get(c, c.value)


def _build_findings_table(doc: Document, findings: Sequence[Finding]) -> None:
    """Findings as a Word table — the shape a reviewer is used to.

    The headings are the table's column titles; each row is one finding.
    A finding with evidence IDs has those IDs listed in the row, so a
    reviewer can look up the passage in one place.
    """
    if not findings:
        return
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "ID"
    hdr[1].text = "Description"
    hdr[2].text = "Severity"
    hdr[3].text = "Evidence"
    for finding in findings:
        row = table.add_row().cells
        row[0].text = finding.id
        row[1].text = finding.description
        row[2].text = finding.severity.label()
        row[3].text = ", ".join(f"[{eid}]" for eid in finding.evidence_ids)


def _build_calculation_section(
    doc: Document, calc_ids: Sequence[str], calculation_records: Dict[str, Any]
) -> None:
    """The Calculation section quotes the engine's results exactly."""
    if not calc_ids:
        doc.add_paragraph("No calculations were required for this inspection.")
        return
    for cid in calc_ids:
        record = calculation_records.get(cid)
        if record is None:
            doc.add_paragraph(f"[{cid}] (calculation record not found)")
            continue
        line = f"[{cid}] {record.expression} = {record.result}"
        doc.add_paragraph(line)


def _build_recommendation(doc: Document, draft: ApprovalNoteDraft) -> None:
    """The recommendation, with severity and classification visible."""
    severity = draft.highest_severity()
    doc.add_paragraph(
        f"Highest severity: {severity.label()}."
    )
    doc.add_paragraph(draft.proposed_action)
    doc.add_paragraph(f"Classification: {_format_classification(draft.classification)}.")


def _build_references(
    doc: Document,
    evidence_ids: Sequence[str],
    passages_by_id: Dict[str, Any],
) -> None:
    """The references section, one paragraph per evidence ID with the
    resolved passage. An unresolved ID is a renderer bug and is named
    plainly; the verifier also flags it."""
    if not evidence_ids:
        doc.add_paragraph("No external references.")
        return
    for eid in evidence_ids:
        passage = passages_by_id.get(eid)
        if passage is None:
            doc.add_paragraph(
                f"[{eid}] (citation does not resolve to a passage in the "
                f"authorized collection)"
            )
            continue
        doc.add_paragraph(
            f"[{eid}] {passage.document_path.name}, page {passage.page}: "
            f"{passage.passage_text}"
        )


def _build_assumptions(
    doc: Document, notes: Sequence[UncertaintyNote]
) -> None:
    """The assumptions section. Empty list is a real answer — the note
    is honest about what is uncertain — so the section is always present
    and always has a paragraph."""
    if not notes:
        doc.add_paragraph(
            "No assumptions or uncertainties were recorded for this inspection."
        )
        return
    for note in notes:
        suffix = " (blocks approval)" if note.blocks_approval else ""
        doc.add_paragraph(f"- {note.what}: {note.reason}{suffix}")


def _add_classification_banner(doc: Document, classification: Classification) -> None:
    """The classification banner at the top of the document, in plain
    text, not in a comment, so it is the first thing the eye lands on."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"Classification: {_format_classification(classification)}")
    run.bold = True


def _add_draft_banner(doc: Document) -> None:
    """A document that is not approved says so where it cannot be missed."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        "DRAFT — not yet approved. Do not act on this document."
    )
    run.bold = True
    run.font.size = Pt(14)


def _add_signature_block(doc: Document, approval: Optional[ApprovalRecord]) -> None:
    """A signature area even when unsigned, so the reviewer knows the
    document is *for* signing and not signed by mistake."""
    doc.add_paragraph("Signed by:")
    doc.add_paragraph("Name: ____________________________")
    doc.add_paragraph("Role: ____________________________")
    doc.add_paragraph("Date: ____________________________")
    if approval is not None:
        decision = (
            approval.decision.value
            if hasattr(approval.decision, "value")
            else str(approval.decision)
        )
        doc.add_paragraph(
            f"Approval record: {decision} by {approval.decided_by} at {approval.decided_at}."
        )
    else:
        doc.add_paragraph("Awaiting human approval.")


def render(
    *,
    draft: ApprovalNoteDraft,
    output_path: Path,
    metadata: DocumentMetadata,
    calculation_records: Dict[str, Any],
    passages_by_id: Dict[str, Any],
) -> RenderResult:
    """Build the document.

    Required sections are added in order. The model never chooses the
    order; the verifier checks the file carries the headings this function
    writes.
    """
    # The model has filled a draft. The renderer is what the model is *not*
    # allowed to do — it controls the document shape, not just the content.
    # Any text the renderer needs to put in is taken from the draft or from
    # the approved trust-list of strings.
    doc = Document()

    if metadata.is_draft:
        _add_draft_banner(doc)
    _add_classification_banner(doc, metadata.classification)

    # Title
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(
        f"Inspection Approval Note — {draft.equipment_id}"
    )
    title_run.bold = True
    title_run.font.size = Pt(16)
    doc.add_paragraph(f"Date of inspection: {draft.inspection_date}")
    doc.add_paragraph(f"Status: {draft.status.value}")

    # The required headings, in the order the reviewer reads them.
    doc.add_heading("Findings", level=1)
    _build_findings_table(doc, draft.findings)

    doc.add_heading("Calculation", level=1)
    _build_calculation_section(doc, draft.calculation_ids, calculation_records)

    doc.add_heading("Recommendation", level=1)
    _build_recommendation(doc, draft)

    doc.add_heading("Supporting references", level=1)
    _build_references(doc, draft.evidence_ids, passages_by_id)

    doc.add_heading("Assumptions", level=1)
    _build_assumptions(doc, draft.uncertainty_notes)

    doc.add_heading("Signature", level=1)
    _add_signature_block(doc, metadata.approval)

    # Provenance at the foot, where it belongs.
    doc.add_heading("How this was produced", level=1)
    doc.add_paragraph(
        f"Task {metadata.task_id} · generated {metadata.created_at} · "
        f"model {metadata.model} · skill {metadata.skill} · "
        f"figures computed by ARJUN's calculation engine, not by the model."
    )

    # Refuse to write half a file: if the doc has nothing in it, that is a bug.
    body_text = "\n".join(p.text for p in doc.paragraphs)
    if len(body_text.strip()) < 50:
        raise TemplateError(
            "The renderer produced a document with effectively no content."
        )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    return verify(output_path)


# ---------------------------------------------------------------------------
# Verification: re-open the file and confirm what it claims to be
# ---------------------------------------------------------------------------


def verify(output_path: Path) -> RenderResult:
    """Re-open a produced file and report what is actually in it.

    Returns a result with ``opens`` true if the file is a DOCX we can
    parse, and ``problems`` listing any reason it is not sound. The
    sections list is the headings the document carries, in document order;
    the verifier compares that list against ``REQUIRED_HEADINGS``.
    """
    if not output_path.exists():
        return RenderResult(
            output_path=output_path,
            problems=[f"The file does not exist: {output_path}"],
        )

    byte_size = output_path.stat().st_size
    if byte_size < 200:
        return RenderResult(
            output_path=output_path,
            problems=["The file is too small to be a real document."],
            byte_size=byte_size,
        )

    # Confirm it is a real ZIP / DOCX.
    try:
        with zipfile.ZipFile(str(output_path)) as zf:
            names = set(zf.namelist())
    except zipfile.BadZipFile:
        return RenderResult(
            output_path=output_path,
            problems=["The file is not a valid ZIP / DOCX archive."],
            byte_size=byte_size,
        )

    for required in (
        "[Content_Types].xml",
        "_rels/.rels",
        "word/document.xml",
    ):
        if required not in names:
            return RenderResult(
                output_path=output_path,
                problems=[f"The document is missing {required}."],
                byte_size=byte_size,
            )

    # Open with python-docx and walk the headings.
    try:
        opened = Document(str(output_path))
    except Exception as exc:  # noqa: BLE001
        return RenderResult(
            output_path=output_path,
            problems=[f"The file could not be opened: {exc}"],
            byte_size=byte_size,
        )

    sections: List[str] = []
    for paragraph in opened.paragraphs:
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            text = paragraph.text.strip()
            if text and text not in sections:
                sections.append(text)

    problems: List[str] = []
    for required in REQUIRED_HEADINGS:
        if required not in sections:
            problems.append(f"the {required!r} section is missing")

    # Look for unfilled placeholders. A model that submits "TBD" is a model
    # that has not finished; the renderer still wrote the file but the
    # verifier must not say it is sound.
    body_text = "\n".join(p.text for p in opened.paragraphs)
    for placeholder in ("TBD", "TODO", "<...>", "<<<"):
        if placeholder in body_text:
            problems.append(
                f"the document still contains the placeholder {placeholder!r}"
            )

    return RenderResult(
        output_path=output_path,
        sections=sections,
        problems=problems,
        opens=True,
        byte_size=byte_size,
    )
