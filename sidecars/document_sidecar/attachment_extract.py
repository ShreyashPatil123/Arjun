#!/usr/bin/env python3
"""Turn one attached file into something a local model can actually read.

Why this exists
---------------
Attachments used to go straight to the vision model, so anything that was not
a PNG, JPEG or WebP was refused — including PDF, which is the format people
actually attach. Forcing a PDF through OCR would also be wrong: most PDFs
carry a text layer, and rendering pages to images to read text that is already
present is slower, lossier, and needlessly uses the GPU.

So this routes by what the file *is*:

    text-native (txt, md, csv, json)  ->  decoded, never OCR'd
    pdf                               ->  page by page: a page with a text
                                          layer is read with pypdf, a page
                                          without one is rendered to PNG for
                                          the OCR model, and a page that is
                                          neither is reported unread
    docx                              ->  paragraphs and tables, never OCR'd
    xlsx                              ->  sheet cells, never OCR'd
    pptx                              ->  slide text and speaker notes, never OCR'd
    anything else                     ->  refused by name, with the reason

Only a page with no text layer reaches the OCR model. That is the point — and
it is decided per page, because a PDF is not obliged to be all one thing.

Dependencies
------------
Nothing new. `pypdf`, `fitz` (PyMuPDF) and `docx` are already installed for
the document sidecar. PyMuPDF is what renders pages — a self-contained wheel,
so there is no Poppler or Ghostscript to install and nothing reaches the
network.

Contract
--------
    python attachment_extract.py <path> <out_dir> [--max-pages N]

Prints one JSON object on stdout:

    {"kind": "pdf-text"|"pdf-scan"|"pdf-mixed"|"text"|"docx"|"xlsx"|"pptx",
     "text": "...",          # the text layer alone; empty for pdf-scan
     "pages": 6,             # real page/sheet count, or 1
     "pageImages": [...],    # absolute PNG paths of the pages needing OCR
     "pageDetail": [...],    # PDFs only: one entry per page, in page order
     "unreadPages": [4, 5],  # pages nothing could be made of
     "truncated": false}     # true whenever unreadPages is non-empty

`pageDetail` is the authoritative account of a PDF, one entry per page:

    {"page": 1, "source": "text",   "text": "..."}
    {"page": 2, "source": "ocr",    "image": "/abs/page-2.png", "layerText": ""}
    {"page": 3, "source": "unread", "why": "..."}

Every page of the document appears exactly once. A page that could not be read
says so; it is never left out, because a missing entry and a blank page are
different facts and the caller cannot tell them apart afterwards.

or, on refusal:

    {"error": "..."}

Every number reported is counted, never estimated: the caller shows them to
the person as "6 pages", and a guess there would be a lie.
"""

import json
import os
import sys
import zipfile

# A page is rendered at this width when it has to go to the OCR model. The
# model normalises coordinates against whatever it is given, and its encoder
# works at roughly this scale; larger costs time for no more detail.
RENDER_WIDTH = 1000

# Rendering is the expensive path, so it is bounded. A hundred-page scan is a
# batch job, not a chat attachment, and saying so beats a silent long wait.
DEFAULT_MAX_PAGES = 12

# Enough of a text layer to call it a text PDF. Below this the "text layer" is
# usually a header or a stray watermark on an otherwise scanned page.
MIN_TEXT_CHARS_PER_PAGE = 24

TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".csv", ".json", ".log", ".tsv"}


def fail(message):
    print(json.dumps({"error": message}))
    return 0


def read_text_native(path):
    """Decode a text file without guessing wildly at its encoding.

    UTF-8 first because that is what these files almost always are; then
    UTF-16, which is what Notepad and Excel still emit; then latin-1, which
    cannot fail and at least preserves byte structure.
    """
    raw = open(path, "rb").read()
    for encoding in ("utf-8-sig", "utf-16", "latin-1"):
        try:
            return raw.decode(encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return raw.decode("latin-1", "replace")


def extract_pdf(path, out_dir, max_pages):
    """Read a PDF, deciding page by page whether it needs the OCR model.

    Why the decision is per page
    ---------------------------
    This used to compare the *whole file's* text against one threshold and
    label the result `pdf-text` or `pdf-scan`. Real documents are not one or
    the other. A digitally-produced cover sheet in front of scanned drawings is
    the ordinary shape of an engineering document, and against an aggregate
    threshold the cover alone carried the whole file over the line: the scanned
    pages were never rendered, never OCR'd, and never mentioned. The reader got
    the cover and no sign that anything else existed.

    The threshold was capped at three pages' worth (`min(total, 3)`), so it
    never grew with the document. One readable page anywhere in a hundred-page
    scan was enough to declare the whole thing text.

    Both failures were silent, which is what made them serious. `truncated`
    came back `False`, so nothing downstream could tell a PDF that had been
    read from one that had been skipped.

    What comes back
    ---------------
    `pageDetail` is the authoritative record: one entry per page, in page
    order, saying how that page was handled. Every page appears in it exactly
    once, including the ones nothing could be done with — a page this function
    cannot read is reported as unread, never omitted.
    """
    import pypdf

    reader = pypdf.PdfReader(path)
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception:
            return {"error": "This PDF is password protected."}

    total = len(reader.pages)
    if total == 0:
        return {"error": "This PDF has no pages in it."}

    # What the text layer gives for each page, kept per page rather than
    # concatenated. The concatenation is precisely what lost the page
    # boundaries, and with them any way to tell which pages were covered.
    layer = []
    for page in reader.pages:
        try:
            layer.append((page.extract_text() or "").strip())
        except Exception:
            # One unreadable page is not an unreadable document. It falls to
            # OCR below like any other page with no usable text.
            layer.append("")

    # The per-page decision. The same threshold as before, applied to the page
    # it is about rather than to a sum across pages.
    needs_ocr = [i for i in range(total) if len(layer[i]) < MIN_TEXT_CHARS_PER_PAGE]

    detail = [None] * total
    for index in range(total):
        if index not in needs_ocr:
            detail[index] = {"page": index + 1, "source": "text",
                             "text": layer[index]}

    rendered = []
    if needs_ocr:
        import fitz

        doc = fitz.open(path)
        try:
            for index in needs_ocr:
                # Rendering is the expensive path, so it stays bounded — but a
                # page dropped for the budget is now *recorded* as dropped
                # rather than vanishing.
                if len(rendered) >= max_pages:
                    detail[index] = {
                        "page": index + 1, "source": "unread",
                        "why": "the limit of %d rendered pages was reached" % max_pages,
                    }
                    continue
                try:
                    page = doc.load_page(index)
                    rect = page.rect
                    scale = RENDER_WIDTH / rect.width if rect.width else 1.0
                    pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
                    target = os.path.join(out_dir, "page-%d.png" % (index + 1))
                    pix.save(target)
                except Exception as error:  # noqa: BLE001 - shown to a person
                    detail[index] = {"page": index + 1, "source": "unread",
                                     "why": "the page could not be rendered: %s" % error}
                    continue
                rendered.append(target)
                detail[index] = {
                    "page": index + 1, "source": "ocr", "image": target,
                    # Whatever the text layer did hold, below the threshold.
                    # Kept as a fallback for when OCR comes back with nothing: a
                    # page carrying "Figure 3 continued" over an unreadable scan
                    # should report that much rather than nothing at all.
                    "layerText": layer[index],
                }
        finally:
            doc.close()

    unread = [d["page"] for d in detail if d["source"] == "unread"]

    # Nothing readable at all is a failure, not a result — the rule
    # `extract_pptx` already follows. A document of nothing but "could not
    # read" lets a summary be written about a file that was never opened.
    if not rendered and not any(d["source"] == "text" for d in detail):
        return {"error": "no page of this PDF could be read; the file may be corrupt"}

    if not needs_ocr:
        kind = "pdf-text"
    elif len(needs_ocr) == total:
        kind = "pdf-scan"
    else:
        kind = "pdf-mixed"

    return {
        "kind": kind,
        # The digital text alone, as before, for any reader that wants one
        # blob. Merging this with the OCR output in page order is the caller's
        # job, because only the caller has the OCR output.
        "text": "\n\n".join(d["text"] for d in detail if d["source"] == "text"),
        "pages": total,
        # The images from `pageDetail`, in page order. Its own field because
        # "is there OCR work?" is the caller's first question, and this answers
        # it without walking the detail.
        "pageImages": rendered,
        "pageDetail": detail,
        "unreadPages": unread,
        "truncated": bool(unread),
    }


def extract_docx(path):
    import docx

    document = docx.Document(path)
    blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    return {"kind": "docx", "text": "\n\n".join(blocks), "pages": 1,
            "pageImages": [], "truncated": False}


SHEET_NS = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"
RELS_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"


def column_index(reference):
    """`A1` -> 0, `B3` -> 1, `AA7` -> 26. `None` when there is no reference.

    A cell carries the square it sits in, and until this was read the reader
    laid cells out by their order in the XML instead. Those are not the same
    thing: a writer is free to omit an empty cell entirely — ARJUN's own does,
    emitting nothing at all for `Cell::Empty` — and every cell after the gap
    then shifts one column left. A value under the wrong heading is worse than
    a value that is missing, because nothing about it looks wrong.
    """
    letters = ""
    for character in reference:
        if character.isalpha():
            letters += character
        else:
            break
    if not letters:
        return None
    index = 0
    for character in letters.upper():
        index = index * 26 + (ord(character) - ord("A") + 1)
    return index - 1


def cell_value(cell, shared):
    """The text of one cell, whichever of the several ways it is stored.

    OOXML has more than one place to keep a value, and the reader used to look
    in exactly one of them — `<v>`, interpreted as a shared-string index when
    `t="s"`. That covers what Excel writes and misses what ARJUN writes: this
    codebase's own workbook writer emits every label as an *inline* string,
    `<c t="inlineStr"><is><t>...</t></is></c>`, which has no `<v>` element at
    all. So a row reading `Inspection result | 9.0` came back as ` | 9.0` — the
    number kept, the word for it dropped, and no sign anything had gone.
    """
    kind = cell.get("t")

    # Inline: the text lives in the cell, and there is no <v> to find.
    if kind == "inlineStr":
        inline = cell.find("%sis" % SHEET_NS)
        if inline is None:
            return ""
        return "".join(t.text or "" for t in inline.iter("%st" % SHEET_NS))

    value = cell.find("%sv" % SHEET_NS)
    raw = "" if value is None or value.text is None else value.text

    # Shared: <v> is an index into the workbook-wide string table.
    if kind == "s":
        try:
            index = int(raw)
        except ValueError:
            return ""
        return shared[index] if 0 <= index < len(shared) else ""

    if kind == "b":
        return "TRUE" if raw.strip() == "1" else "FALSE"

    # An Excel error (#DIV/0!, #REF!) is reported as it stands. Blanking it
    # would turn a broken figure into a missing one, and a reader deciding
    # whether to trust a workbook needs to see it.
    if kind == "e":
        return raw

    formula = cell.find("%sf" % SHEET_NS)
    if formula is not None:
        expression = (formula.text or "").strip()
        if expression and raw:
            # Both, because they answer different questions. The cached value
            # is what the workbook says; the formula is how it got there, and
            # a calculation workbook exists to show exactly that.
            return "%s (=%s)" % (raw, expression)
        if expression:
            # A formula Excel has never evaluated has no cache. Saying so beats
            # reporting the cell as empty.
            return "=%s" % expression
    return raw


def sheet_parts(archive):
    """Every sheet as `(name, part)`, in the order the workbook lists them.

    Read through `workbook.xml` and its relationships rather than by globbing
    the parts directory, for two reasons. The names are only in the workbook —
    "Calculation" is a great deal more use to a reader than "sheet1" — and the
    order is only correct there: sorting part names as text puts `sheet10`
    ahead of `sheet2`.
    """
    from xml.etree import ElementTree as ET

    names = archive.namelist()
    targets = {}
    if "xl/_rels/workbook.xml.rels" in names:
        root = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
        for relationship in root:
            target = (relationship.get("Target") or "").lstrip("/")
            if not target:
                continue
            if not target.startswith("xl/"):
                target = "xl/" + target
            targets[relationship.get("Id")] = target

    named = []
    if "xl/workbook.xml" in names:
        root = ET.fromstring(archive.read("xl/workbook.xml"))
        for sheet in root.iter("%ssheet" % SHEET_NS):
            part = targets.get(sheet.get("%sid" % RELS_NS))
            if part and part in names:
                named.append((sheet.get("name") or "Sheet", part))
    if named:
        return named

    # No usable workbook part. Fall back to the sheet files themselves, ordered
    # by their number rather than as text.
    def number(part):
        digits = "".join(c for c in os.path.basename(part) if c.isdigit())
        return int(digits) if digits else 0

    parts = [n for n in names
             if n.startswith("xl/worksheets/sheet") and n.endswith(".xml")]
    return [(os.path.splitext(os.path.basename(p))[0], p)
            for p in sorted(parts, key=number)]


def extract_xlsx(path, max_rows=2000):
    """Read sheet values straight from the package.

    `openpyxl` is not installed, and adding it for one format would be a new
    dependency in an air-gapped deployment. An .xlsx is a zip of XML, and the
    standard library reads both halves of it.

    The rule this follows is that a cell's *position* and a cell's *label* are
    part of its value. A workbook read as a bag of numbers with the headings
    dropped and the columns shifted is not a workbook that was read.
    """
    from xml.etree import ElementTree as ET

    with zipfile.ZipFile(path) as archive:
        shared = []
        if "xl/sharedStrings.xml" in archive.namelist():
            root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
            for item in root.findall("%ssi" % SHEET_NS):
                shared.append("".join(t.text or "" for t in item.iter("%st" % SHEET_NS)))

        sheets = sheet_parts(archive)
        out, count, truncated = [], 0, False

        for name, part in sheets:
            root = ET.fromstring(archive.read(part))
            rows = []
            for row in root.iter("%srow" % SHEET_NS):
                if count >= max_rows:
                    truncated = True
                    break
                # Keyed on the column the cell says it is in, so a gap stays a
                # gap instead of pulling everything after it one place left.
                cells, widest = {}, -1
                for cell in row.findall("%sc" % SHEET_NS):
                    index = column_index(cell.get("r") or "")
                    if index is None:
                        index = widest + 1
                    cells[index] = cell_value(cell, shared)
                    widest = max(widest, index)
                values = [cells.get(i, "") for i in range(widest + 1)]
                if any(v.strip() for v in values):
                    rows.append(" | ".join(values))
                count += 1

            if rows:
                # Which sheet a row came from, said once per sheet — but only
                # when there is more than one sheet to be confused about. A
                # figure quoted out of a five-tab workbook with no idea which
                # tab it came from cannot be checked; a header over the only
                # sheet there is answers a question nobody can ask, and every
                # single-sheet attachment would carry the noise.
                #
                # Keyed on how many sheets the *workbook* declares, not on how
                # many turned out to have rows: a three-sheet workbook with
                # content on one of them is exactly the case where naming it
                # earns its line.
                if len(sheets) > 1:
                    out.append("--- sheet: %s ---" % name)
                out.extend(rows)
            if truncated:
                break

        return {"kind": "xlsx", "text": "\n".join(out), "pages": len(sheets),
                "pageImages": [], "truncated": truncated}


def extract_pptx(path, max_slides=500):
    """Read slide text straight from the package.

    Same reasoning as `extract_xlsx`: `python-pptx` is not installed and adding
    it for one format would be a new dependency in an air-gapped deployment. A
    .pptx is a zip of XML — each slide is `ppt/slides/slideN.xml` and every run
    of text in it is an `<a:t>` element, which the standard library reads.

    Slides are emitted in numeric order and separated by a heading, because a
    reviewer asking "what does slide 7 say" needs the boundary to survive into
    the text. Speaker notes are included: on an approval deck the note under
    the slide is often where the actual commitment is written.
    """
    from xml.etree import ElementTree as ET

    drawing = "{http://schemas.openxmlformats.org/drawingml/2006/main}"

    def slide_number(name):
        digits = "".join(c for c in os.path.basename(name) if c.isdigit())
        return int(digits) if digits else 0

    def text_of(member, archive):
        # A slide that will not parse costs that slide, not the deck. One
        # corrupt part in a 60-slide package used to raise out of the whole
        # extraction, so a reviewer got "could not be read" for a file that
        # was 59/60ths readable -- and no way to tell which slide was bad.
        try:
            root = ET.fromstring(archive.read(member))
        except (ET.ParseError, KeyError):
            return None
        runs = [t.text for t in root.iter("%st" % drawing) if t.text]
        return "\n".join(run.strip() for run in runs if run.strip())

    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        slides = sorted(
            (n for n in names
             if n.startswith("ppt/slides/slide") and n.endswith(".xml")),
            key=slide_number,
        )
        notes = {
            slide_number(n): n
            for n in names
            if n.startswith("ppt/notesSlides/notesSlide") and n.endswith(".xml")
        }

        blocks, truncated, unreadable = [], False, 0
        for index, member in enumerate(slides, start=1):
            if index > max_slides:
                truncated = True
                break
            body = text_of(member, z)
            note = notes.get(slide_number(member))
            note_text = text_of(note, z) if note else ""
            section = ["Slide %d" % index]
            if body is None:
                # Named rather than skipped silently. A gap a reader can
                # see is recoverable; one they cannot is a summary with a
                # hole in it.
                section.append("[this slide could not be read]")
                unreadable += 1
            elif body:
                section.append(body)
            if note_text:
                section.append("Speaker notes: %s" % note_text)
            # A slide with a diagram and no text still gets its heading, so the
            # numbering stays aligned with the deck a person is looking at.
            blocks.append("\n".join(section))

        # Every slide unreadable is a failure, not a result. Returning a
        # document of nothing but placeholders would let a summary be
        # written about a deck that was never read.
        if slides and unreadable == len(blocks):
            return {"error": "no slide in this deck could be read; the file may be corrupt"}

        return {"kind": "pptx", "text": "\n\n".join(blocks),
                "pages": len(slides), "pageImages": [], "truncated": truncated}


def main():
    if len(sys.argv) < 3:
        return fail("usage: attachment_extract.py <path> <out_dir> [--max-pages N]")
    path, out_dir = sys.argv[1], sys.argv[2]
    max_pages = DEFAULT_MAX_PAGES
    if "--max-pages" in sys.argv:
        try:
            max_pages = int(sys.argv[sys.argv.index("--max-pages") + 1])
        except (IndexError, ValueError):
            return fail("--max-pages needs a number")

    if not os.path.isfile(path):
        return fail("that file is not there any more")
    os.makedirs(out_dir, exist_ok=True)
    suffix = os.path.splitext(path)[1].lower()

    try:
        if suffix in TEXT_SUFFIXES:
            result = {"kind": "text", "text": read_text_native(path), "pages": 1,
                      "pageImages": [], "truncated": False}
        elif suffix == ".pdf":
            result = extract_pdf(path, out_dir, max_pages)
        elif suffix == ".docx":
            result = extract_docx(path)
        elif suffix == ".xlsx":
            result = extract_xlsx(path)
        elif suffix == ".pptx":
            result = extract_pptx(path)
        else:
            return fail("%s is not a document ARJUN can read."
                        % (suffix or "this file type"))
    except Exception as error:  # noqa: BLE001 - the caller shows this to a person
        return fail("%s could not be read: %s" % (os.path.basename(path), error))

    if "error" in result:
        return fail(result["error"])
    print(json.dumps(result))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
